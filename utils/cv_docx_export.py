"""
Convert optimized CV markdown to a professionally formatted Word document (.docx).

Used when LibreOffice is unavailable — produces real Word list styles and typography
closer to a hand-crafted résumé than ODT via odfpy.
"""

from __future__ import annotations

import html
import io
import re
from typing import Optional

from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.shared import Inches, Pt, RGBColor

_BULLET_PREFIX_RE = re.compile(r"^[\-\u2022\u25AA\u25E6]\s*(.+)$")
_ASTERISK_BULLET_RE = re.compile(r"^\*\s+(.+)$")
_CONTACT_RE = re.compile(r"[^\s@]+@[^\s@]+\.[^\s@]+")
_SKILL_CATEGORY_RE = re.compile(
    r"^[\-\u2022\u25AA\*]?\s*\*{0,2}([^*\n:]{3,}):\*{0,2}\s*(.+)$"
)
_ROLE_SPLIT_RE = re.compile(r"^(.+?)\s[—–-]\s(.+)$")

_SECTION_ALIASES = {
    "work experience": "Relevant Experience",
    "skills": "Core Technologies",
    "professional summary": "",  # no heading — summary flows under title like source résumé
}

_DEGREE_HINT_RE = re.compile(
    r"\b(bachelor|master|doctor|ph\.?d|mba|llb|ba|bs|b\.s|m\.s|associate|diploma|certificate)\b",
    re.IGNORECASE,
)


def _decode_text(text: str) -> str:
    prev = None
    current = text or ""
    while prev != current:
        prev = current
        current = html.unescape(current)
    return current.strip()


def _strip_wrapping_asterisks(text: str) -> str:
    stripped = text.strip()
    if stripped.startswith("**") and stripped.endswith("**") and len(stripped) > 4:
        return stripped[2:-2].strip()
    if (
        stripped.startswith("*")
        and stripped.endswith("*")
        and not stripped.startswith("**")
        and len(stripped) > 2
    ):
        return stripped[1:-1].strip()
    return stripped


def _is_contact_line(text: str) -> bool:
    return bool(_CONTACT_RE.search(text))


def _is_date_line(text: str) -> bool:
    stripped = text.strip()
    return (
        stripped.startswith("*")
        and stripped.endswith("*")
        and not stripped.startswith("**")
        and any(ch.isdigit() for ch in stripped)
    )


def _parse_bullet_line(text: str) -> Optional[str]:
    stripped = text.strip()
    if _is_date_line(stripped):
        return None
    match = _BULLET_PREFIX_RE.match(stripped)
    if match:
        return _decode_text(match.group(1))
    match = _ASTERISK_BULLET_RE.match(stripped)
    if match:
        return _decode_text(match.group(1))
    return None


def _parse_skill_line(text: str) -> Optional[tuple[str, str]]:
    stripped = text.strip()
    match = _SKILL_CATEGORY_RE.match(stripped)
    if match:
        return _decode_text(match.group(1).rstrip(":")), _decode_text(match.group(2))
    if ":" in stripped and not stripped.startswith("#"):
        label, _, body = stripped.partition(":")
        label = _decode_text(label.lstrip("-•▪* ").strip())
        body = _decode_text(body.strip())
        if len(label) >= 3 and body:
            return label, body
    return None


def _set_run_font(
    run,
    *,
    size_pt: float,
    bold: bool = False,
    italic: bool = False,
    color: Optional[RGBColor] = None,
) -> None:
    run.font.name = "Arial"
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def _add_paragraph(
    doc: Document,
    text: str,
    *,
    size_pt: float = 10.5,
    bold: bool = False,
    italic: bool = False,
    color: Optional[RGBColor] = None,
    space_before_pt: float = 0,
    space_after_pt: float = 2,
    line_spacing: float = 1.15,
) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(space_before_pt)
    paragraph.paragraph_format.space_after = Pt(space_after_pt)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    paragraph.paragraph_format.line_spacing = line_spacing
    run = paragraph.add_run(_decode_text(text))
    _set_run_font(run, size_pt=size_pt, bold=bold, italic=italic, color=color)


def _add_section_heading(doc: Document, title: str) -> None:
    section_title = _SECTION_ALIASES.get(title.lower(), title)
    if not section_title:
        return
    _add_paragraph(
        doc,
        section_title,
        size_pt=11,
        bold=True,
        space_before_pt=10,
        space_after_pt=4,
    )


def _normalize_skill_separators(body: str) -> str:
    """Use middle dots between skills like the source résumé."""
    if " • " in body:
        return body
    if "," in body:
        parts = [_decode_text(p.strip()) for p in body.split(",") if p.strip()]
        if len(parts) > 1:
            return " • ".join(parts)
    return body


def _add_skill_line(doc: Document, label: str, body: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.paragraph_format.line_spacing = 1.1
    label_run = paragraph.add_run(f"{label}: ")
    _set_run_font(label_run, size_pt=10.5, bold=True)
    body_run = paragraph.add_run(_normalize_skill_separators(body))
    _set_run_font(body_run, size_pt=10.5, bold=False)


def _add_bullet(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(1)
    paragraph.paragraph_format.line_spacing = 1.12
    run = paragraph.add_run(_decode_text(text))
    _set_run_font(run, size_pt=10.5)


def _add_role_block(doc: Document, line: str) -> None:
    """Split ``Company — Title`` into title + company lines like the source résumé."""
    match = _ROLE_SPLIT_RE.match(line.strip())
    if match:
        company = _decode_text(match.group(1))
        title = _decode_text(match.group(2))
        _add_paragraph(doc, title, size_pt=10.5, bold=True, space_before_pt=6, space_after_pt=0)
        _add_paragraph(doc, company, size_pt=10.5, bold=True, space_after_pt=2)
    else:
        _add_paragraph(doc, line, size_pt=10.5, bold=True, space_before_pt=6, space_after_pt=2)


def _add_education_block(doc: Document, line: str) -> None:
    """Render education as institution line then degree line(s), matching classic résumés."""
    match = _ROLE_SPLIT_RE.match(line.strip())
    if match:
        left = _decode_text(match.group(1))
        right = _decode_text(match.group(2))
        if _DEGREE_HINT_RE.search(left):
            degree, institution = left, right
        elif _DEGREE_HINT_RE.search(right):
            degree, institution = right, left
        else:
            degree, institution = left, right
        _add_paragraph(doc, institution, size_pt=10.5, bold=True, space_before_pt=6, space_after_pt=0)
        _add_paragraph(doc, degree, size_pt=10.5, bold=True, space_after_pt=2)
    else:
        _add_paragraph(doc, line, size_pt=10.5, bold=True, space_before_pt=6, space_after_pt=2)


def markdown_cv_to_docx_bytes(cv_markdown: str) -> bytes:
    """
    Convert optimized CV markdown to a Word ``.docx`` file.

    Args:
        cv_markdown: Markdown-formatted CV from the optimizer

    Returns:
        DOCX file contents as bytes
    """
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    lines = (cv_markdown or "").replace("\r\n", "\n").split("\n")
    current_section = ""
    in_summary = False
    pending_title: Optional[str] = None
    pending_contact: Optional[str] = None

    def _flush_header_block() -> None:
        nonlocal pending_title, pending_contact
        if pending_contact:
            _add_paragraph(
                doc,
                pending_contact,
                size_pt=9,
                color=RGBColor(0x66, 0x66, 0x66),
                space_after_pt=4,
            )
            pending_contact = None
        if pending_title:
            _add_paragraph(doc, pending_title, size_pt=11, bold=True, space_after_pt=6)
            pending_title = None

    for raw_line in lines:
        line = raw_line.rstrip()
        if not line:
            continue
        stripped = line.strip()

        if line.startswith("# "):
            name = _decode_text(line[2:].strip())
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(2)
            run = paragraph.add_run(name)
            _set_run_font(run, size_pt=14, bold=False)
            continue

        if line.startswith("## "):
            _flush_header_block()
            current_section = _decode_text(line[3:].strip())
            in_summary = current_section.lower() == "professional summary"
            _add_section_heading(doc, current_section)
            continue

        if not current_section:
            if _is_contact_line(stripped):
                pending_contact = stripped
                continue
            if stripped.startswith("**") or (
                stripped.startswith("*") and stripped.endswith("*") and not _is_date_line(stripped)
            ):
                subtitle = _strip_wrapping_asterisks(stripped)
                if subtitle and not any(ch.isdigit() for ch in subtitle):
                    pending_title = subtitle
                    continue

        if line.startswith("### "):
            block_line = line[4:].strip()
            if current_section.lower() == "education":
                _add_education_block(doc, block_line)
            else:
                _add_role_block(doc, block_line)
            continue

        if _is_date_line(stripped):
            continue

        if _is_contact_line(stripped):
            _add_paragraph(
                doc,
                stripped,
                size_pt=9,
                color=RGBColor(0x66, 0x66, 0x66),
                space_after_pt=8,
            )
            continue

        skill = _parse_skill_line(stripped)
        if skill and current_section.lower() in ("skills", "core technologies"):
            _add_skill_line(doc, skill[0], skill[1])
            continue

        bullet = _parse_bullet_line(line)
        if bullet is not None:
            _add_bullet(doc, bullet)
            continue

        if stripped.startswith("**") or (
            stripped.startswith("*") and stripped.endswith("*") and not _is_date_line(stripped)
        ):
            subtitle = _strip_wrapping_asterisks(stripped)
            if subtitle and not any(ch.isdigit() for ch in subtitle):
                _add_paragraph(doc, subtitle, size_pt=11, bold=True, space_after_pt=2)
                continue

        if current_section.lower() in (
            "work experience",
            "relevant experience",
        ) and _ROLE_SPLIT_RE.match(stripped):
            _add_role_block(doc, stripped)
            continue

        if current_section.lower() == "education" and _ROLE_SPLIT_RE.match(stripped):
            _add_education_block(doc, stripped)
            continue

        if in_summary or current_section.lower() == "professional summary":
            _add_paragraph(doc, stripped, size_pt=10.5, space_after_pt=4)
            continue

        if current_section.lower() in ("skills", "core technologies") and "•" in stripped:
            if ":" in stripped:
                label, _, body = stripped.partition(":")
                _add_skill_line(doc, label.strip(" •▪-*"), body.strip())
                continue

        _add_paragraph(doc, stripped, size_pt=10.5)

    _flush_header_block()

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
